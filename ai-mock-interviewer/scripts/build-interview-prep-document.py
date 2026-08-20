import json
import os
import subprocess
import argparse
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_MODULE = os.path.join(ROOT, "public", "interview-prep-data.js")
OUT_PATH = os.path.join(ROOT, "AI-Engineer-Interview-Prep-Curriculum-2026.docx")
DATA_SCIENCE_OUT_PATH = os.path.join(ROOT, "data-science-path", "Data-Science-Interview-Prep-Curriculum-2026.docx")
AI_AGENT_OUT_PATH = os.path.join(ROOT, "ai-agent-engineer-path", "AI-Agent-Engineer-Interview-Prep-Curriculum-2026.docx")

BLUE = "1D4ED8"
DARK = "172554"
INK = "1E293B"
MUTED = "64748B"
PALE = "E8EEF5"
LIGHT = "F8FAFC"
WHITE = "FFFFFF"


def load_curriculum():
    script = "const d=require(process.argv[1]); process.stdout.write(JSON.stringify(d));"
    result = subprocess.run(
        ["node", "-e", script, DATA_MODULE],
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(result.stdout)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_run(run, size=10.5, color=INK, bold=False, italic=False):
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def write_cell(cell, text, *, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    style_run(paragraph.add_run(str(text)), size=size, color=color, bold=bold)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(section, label="AI Engineer Interview Prep · 2026"):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(paragraph.add_run(label), size=8.5, color=MUTED)


def add_role_matrix(doc, matrix):
    doc.add_heading("2026 role priority map", level=1)
    p = doc.add_paragraph("Both roles require strong Python and ML foundations. Their emphasis separates at experimentation, deployment, infrastructure, and operational ownership.")
    p.paragraph_format.space_after = Pt(10)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = [4320, 2520, 2520]
    headers = ["Skill area", "Data Scientist", "ML Engineer"]
    for index, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], PALE)
        write_cell(table.rows[0].cells[index], label, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_header(table.rows[0])
    for item in matrix:
        cells = table.add_row().cells
        write_cell(cells[0], item["skill"])
        write_cell(cells[1], f'{item["dataScientist"]} / 5', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE)
        write_cell(cells[2], f'{item["mlEngineer"]} / 5', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE)
    set_table_geometry(table, widths)


def add_ai_engineer_callout(doc, profile):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    write_cell(cell, "", size=10.5)
    p = cell.paragraphs[0]
    style_run(p.add_run(profile["title"]), size=11, color=DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    style_run(p2.add_run(profile["summary"]), size=10, color=INK)
    set_table_geometry(table, [9360])


def add_domain_index(doc, domains):
    doc.add_heading("Curriculum index", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["No.", "Domain", "Topics"]
    for index, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], PALE)
        write_cell(table.rows[0].cells[index], label, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_header(table.rows[0])
    for index, domain in enumerate(domains, start=1):
        cells = table.add_row().cells
        write_cell(cells[0], f"{index:02}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE)
        write_cell(cells[1], domain["name"], bold=True)
        write_cell(cells[2], len(domain["topics"]), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [720, 7560, 1080])


def add_domains(doc, domains):
    doc.add_page_break()
    doc.add_heading("Complete interview-prep curriculum", level=1)
    for index, domain in enumerate(domains, start=1):
        heading = doc.add_heading(f'{index:02} · {domain["name"]}', level=2)
        heading.paragraph_format.page_break_before = index > 1
        description = doc.add_paragraph(domain["description"])
        description.paragraph_format.space_after = Pt(8)
        description.runs[0].italic = True
        description.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        topics = domain["topics"]
        rows = (len(topics) + 2) // 3
        table = doc.add_table(rows=rows, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for position in range(rows * 3):
            row = position // 3
            col = position % 3
            text = topics[position] if position < len(topics) else ""
            if row % 2:
                set_cell_shading(table.rows[row].cells[col], LIGHT)
            write_cell(table.rows[row].cells[col], text, size=9.2)
        set_table_geometry(table, [3120, 3120, 3120])


def select_data_science_domains(data):
    domain_map = {domain["name"]: domain for domain in data["domains"]}
    stages = data["dataSciencePath"]["stages"]
    names = [name for stage in stages for name in stage["domains"]]
    return [domain_map[name] for name in names if name in domain_map]


def add_stage_index(doc, stages, domains):
    doc.add_heading("Learning path", level=1)
    topic_counts = {domain["name"]: len(domain["topics"]) for domain in domains}
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, label in enumerate(("Stage", "Focus", "Topics")):
        set_cell_shading(table.rows[0].cells[index], PALE)
        write_cell(table.rows[0].cells[index], label, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_header(table.rows[0])
    for index, stage in enumerate(stages, start=1):
        cells = table.add_row().cells
        write_cell(cells[0], index, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE)
        write_cell(cells[1], stage["name"].split("·", 1)[-1].strip(), bold=True)
        write_cell(cells[2], sum(topic_counts.get(name, 0) for name in stage["domains"]), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [900, 7200, 1260])
    for stage in stages:
        heading = doc.add_heading(stage["name"], level=2)
        heading.paragraph_format.keep_with_next = True
        doc.add_paragraph(stage["description"])


def add_topic_stage_index(doc, stages):
    doc.add_heading("Learning path", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, label in enumerate(("Stage", "Focus", "Topics")):
        set_cell_shading(table.rows[0].cells[index], PALE)
        write_cell(table.rows[0].cells[index], label, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER if index != 1 else WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_header(table.rows[0])
    for index, stage in enumerate(stages, start=1):
        cells = table.add_row().cells
        write_cell(cells[0], index, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=BLUE)
        write_cell(cells[1], stage["name"].split("·", 1)[-1].strip(), bold=True)
        write_cell(cells[2], len(stage["topics"]), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [900, 7200, 1260])


def build(track="complete"):
    data = load_curriculum()
    is_data_science = track == "data-science"
    is_ai_agent = track == "ai-agent"
    if is_data_science:
        domains = select_data_science_domains(data)
    elif is_ai_agent:
        domains = [{"name": stage["name"], "description": stage["description"], "topics": stage["topics"]} for stage in data["aiAgentEngineerPath"]["stages"]]
    else:
        domains = data["domains"]
    unique_topics = {topic.lower() for domain in domains for topic in domain["topics"]}
    out_path = AI_AGENT_OUT_PATH if is_ai_agent else (DATA_SCIENCE_OUT_PATH if is_data_science else OUT_PATH)
    title_text = "AI Agent Engineer Interview Prep" if is_ai_agent else ("Data Science Interview Prep" if is_data_science else "AI Engineer Interview Prep")
    subtitle_text = (
        "Job-aligned path for Python, LLM APIs, MCP, RAG, enterprise automation, agent safety, and AWS operations"
        if is_ai_agent else
        "Beginner-to-expert path for statistics, experimentation, machine learning, applied GenAI, and production literacy"
        if is_data_science else
        "Beginner-to-expert curriculum for Data Science, ML Engineering, and production AI"
    )
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    footer_label = "AI Agent Engineer Path · 2026" if is_ai_agent else ("Data Science Interview Path · 2026" if is_data_science else "AI Engineer Interview Prep · 2026")
    add_footer(section, footer_label)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(92)
    title.paragraph_format.space_after = Pt(8)
    style_run(title.add_run(title_text), size=30, color=DARK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(20)
    style_run(subtitle.add_run(subtitle_text), size=14, color=BLUE)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(meta.add_run(f'{len(domains)} domains · {len(unique_topics)} unique topics · Updated {date.today().isoformat()}'), size=10.5, color=MUTED, bold=True)
    purpose = doc.add_paragraph()
    purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
    purpose.paragraph_format.space_before = Pt(24)
    purpose.paragraph_format.left_indent = Inches(0.65)
    purpose.paragraph_format.right_indent = Inches(0.65)
    purpose_text = (
        "A client-facing preparation path for discovering manual workflows, designing governed agents, integrating enterprise tools, deploying on AWS, and proving reliable business outcomes."
        if is_ai_agent else
        "A revision-first path for mastering concepts, implementation, trade-offs, experimentation, business interpretation, and the production context expected in modern Data Science interviews."
        if is_data_science else
        "A revision-first reference for building practical depth across fundamentals, model development, deployment, cloud infrastructure, evaluation, monitoring, system design, and communication."
    )
    style_run(purpose.add_run(purpose_text), size=11, color=INK)

    doc.add_page_break()
    if is_ai_agent:
        add_topic_stage_index(doc, data["aiAgentEngineerPath"]["stages"])
    elif is_data_science:
        add_stage_index(doc, data["dataSciencePath"]["stages"], domains)
    else:
        add_role_matrix(doc, data["roleMatrix"])
        add_ai_engineer_callout(doc, data["aiEngineer"])
    add_domain_index(doc, domains)
    add_domains(doc, domains)
    doc.core_properties.title = f"{title_text} Curriculum 2026"
    doc.core_properties.subject = "Synchronized interview preparation curriculum"
    doc.core_properties.author = "AI Mock Interviewer"
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Created {out_path} with {len(domains)} domains and {len(unique_topics)} unique topics")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--track", choices=("complete", "data-science", "ai-agent"), default="complete")
    args = parser.parse_args()
    build(args.track)
