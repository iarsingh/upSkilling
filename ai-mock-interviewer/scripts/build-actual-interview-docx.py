import os
import re
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SOURCE = os.path.join(ROOT, "actual-interview-questions-and-answers.md")
VERIFIED_ONLY = os.environ.get("VERIFIED_ONLY") == "1"
OUTPUT = os.path.join(ROOT, "Actual-Interview-Questions-and-Answers-Verified.docx" if VERIFIED_ONLY else "Actual-Interview-Questions-and-Answers.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x22, 0x2B, 0x35)
MUTED = RGBColor(0x66, 0x72, 0x7A)
LIGHT_BLUE = "E8EEF5"


def set_font(run, size, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Answer Type" not in doc.styles:
        style = doc.styles.add_style("Answer Type", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Answer Type"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.font.italic = True
    style.font.color.rgb = MUTED
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.keep_with_next = True

    if "Interview Answer" not in doc.styles:
        style = doc.styles.add_style("Interview Answer", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Interview Answer"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK
    style.paragraph_format.left_indent = Inches(0.16)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.25


def parse_source(text):
    categories = []
    current = None
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if line.startswith("## "):
            current = {"name": line[3:].strip(), "questions": []}
            categories.append(current)
        elif line.startswith("### ") and current is not None:
            match = re.match(r"###\s+\d+\.\s+(.+)", line)
            question = match.group(1).strip() if match else line[4:].strip()
            question_type = "Interview question"
            answer_parts = []
            index += 1
            while index < len(lines) and not lines[index].startswith("### ") and not lines[index].startswith("## "):
                value = lines[index].strip()
                if value.startswith("**Type:**"):
                    question_type = value.replace("**Type:**", "", 1).strip()
                elif value == "**Answer:**":
                    pass
                elif value and value != "---" and not re.match(r"^\d+ questions?$", value):
                    answer_parts.append(value)
                index += 1
            current["questions"].append({"question": question, "type": question_type, "answer": "\n".join(answer_parts).strip()})
            continue
        index += 1
    return categories


with open(SOURCE, "r", encoding="utf-8") as handle:
    source_text = handle.read()

categories = parse_source(source_text)
if VERIFIED_ONLY:
    placeholder_pattern = re.compile(
        r"a strong answer should|tailor the response directly to|use the prompt details as acceptance criteria",
        re.IGNORECASE,
    )
    categories = [
        {**category, "questions": [item for item in category["questions"] if not placeholder_pattern.search(item["answer"])]}
        for category in categories
    ]
    categories = [category for category in categories if category["questions"]]
question_count = sum(len(category["questions"]) for category in categories)

doc = Document()
configure_styles(doc)
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.paragraph_format.space_after = Pt(0)
set_font(header.add_run("ACTUAL INTERVIEW HANDBOOK"), 9, MUTED, bold=True)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.paragraph_format.space_before = Pt(0)
add_page_field(footer)
for run in footer.runs:
    set_font(run, 9, MUTED)

# Editorial cover pattern.
for _ in range(5):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)

kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(kicker.add_run("SENIOR ENGINEERING INTERVIEW REFERENCE"), 10, BLUE, bold=True)
kicker.paragraph_format.space_after = Pt(18)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
document_title = "Actual Interview Questions\nand Verified Answers" if VERIFIED_ONLY else "Actual Interview Questions\nand Answers"
set_font(title.add_run(document_title), 28, DARK_BLUE, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(22)
set_font(subtitle.add_run("DevOps · Platform Engineering · SRE · Cloud · MLOps"), 14, MUTED)

summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
summary.paragraph_format.space_after = Pt(8)
set_font(summary.add_run(f"{question_count:,} unique questions · {len(categories)} categories"), 12, INK, bold=True)

generated = doc.add_paragraph()
generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(generated.add_run(f"Prepared {date.today().isoformat()}"), 10, MUTED, italic=True)

doc.add_page_break()

intro = doc.add_heading("How to Use This Handbook", level=1)
doc.add_paragraph(
    "Questions are grouped alphabetically by interview category. Read each question aloud, answer it before reviewing the model response, and adapt behavioral or experience-based material using only facts from your own work."
)

index_heading = doc.add_heading("Category Index", level=1)
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(5.5)
table.columns[1].width = Inches(1.0)
table.rows[0].cells[0].width = Inches(5.5)
table.rows[0].cells[1].width = Inches(1.0)
table.rows[0].cells[0].text = "Category"
table.rows[0].cells[1].text = "Questions"
for cell in table.rows[0].cells:
    set_cell_margins(cell)
    cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
    cell._tc.tcPr[-1].set(qn("w:fill"), LIGHT_BLUE)
    for run in cell.paragraphs[0].runs:
        set_font(run, 10, DARK_BLUE, bold=True)
for category in categories:
    cells = table.add_row().cells
    cells[0].width = Inches(5.5)
    cells[1].width = Inches(1.0)
    cells[0].text = category["name"]
    cells[1].text = str(len(category["questions"]))
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for cell in cells:
        set_cell_margins(cell)
        for run in cell.paragraphs[0].runs:
            set_font(run, 9.5, INK)

doc.add_page_break()

global_number = 1
for category_index, category in enumerate(categories, 1):
    heading = doc.add_heading(f"{category_index}. {category['name']}", level=1)
    count = doc.add_paragraph(f"{len(category['questions'])} question{'s' if len(category['questions']) != 1 else ''}", style="Answer Type")
    for item in category["questions"]:
        q = doc.add_heading(f"Q{global_number}. {item['question']}", level=3)
        q.paragraph_format.keep_with_next = True
        type_para = doc.add_paragraph(f"Type: {item['type']}", style="Answer Type")
        answer = doc.add_paragraph(style="Interview Answer")
        label = answer.add_run("Answer: ")
        set_font(label, 11, DARK_BLUE, bold=True)
        body = answer.add_run(item["answer"] or "Answer pending.")
        set_font(body, 11, INK)
        global_number += 1

doc.core_properties.title = "Actual Interview Questions and Answers"
doc.core_properties.subject = "Senior DevOps, Platform Engineering, SRE, Cloud and MLOps interview handbook"
doc.core_properties.author = "AI Mock Interviewer"
doc.core_properties.keywords = "interview, DevOps, SRE, GCP, AWS, Kubernetes, MLOps"
doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
print(f"Questions: {question_count}; categories: {len(categories)}")
