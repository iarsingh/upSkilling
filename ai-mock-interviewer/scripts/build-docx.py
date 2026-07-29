import json
import re
import os
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_PATH = os.path.join(ROOT, "scripts", "answer-bank", "final-qa-dataset.json")
OUT_PATH = os.path.join(ROOT, "AI-Mock-Interview-Question-Bank.docx")

with open(DATA_PATH, "r") as f:
    entries = json.load(f)

CHEAT_SHEET_PATH = os.path.join(ROOT, "public", "cheat-sheets.json")
with open(CHEAT_SHEET_PATH, "r") as f:
    cheat_sheets = json.load(f)

TOPIC_SECTIONS = [
    "Cloud Platforms & GCP Services",
    "GCP Networking",
    "Kubernetes & Containerization",
    "Infrastructure as Code (IaC)",
    "CI/CD & GitOps",
    "MLOps & AI Platforms",
    "DevSecOps & Cloud Security",
    "Monitoring, Logging & Observability",
    "Programming & Scripting",
    "Databases & Data Services",
    "Messaging & Streaming",
    "Source Control & Artifact Management",
    "Backup & Disaster Recovery",
    "ITSM & Enterprise Tools",
    "Platform Engineering & SRE",
    "Hybrid & Multi-Cloud",
    "Application & API Technologies",
]

def classify_question(entry):
    text = " ".join([
        entry.get("section") or "",
        entry.get("category") or "",
        entry.get("question") or "",
    ]).lower()
    rules = [
        ("GCP Networking", r"\b(vpc|subnet|cidr|cloud router|bgp|vpn|interconnect|cloud nat|dns|private service|network|load balanc|firewall|route|ncc)\b"),
        ("Kubernetes & Containerization", r"\b(kubernetes|gke|container|docker|helm|pod|deployment|statefulset|daemonset|ingress|gateway|hpa|vpa|cluster autoscaler|node pool|taint|toleration|persistent volume)\b"),
        ("Infrastructure as Code (IaC)", r"\b(terraform|infrastructure as code|iac|ansible|remote state|state lock|provider management|drift)\b"),
        ("CI/CD & GitOps", r"\b(ci/cd|cicd|jenkins|pipeline|gitops|argo cd|github actions|gitlab ci|bitbucket pipeline|blue-green|canary|build automation|deployment automation)\b"),
        ("MLOps & AI Platforms", r"\b(mlops|vertex ai|kubeflow|mlflow|model serving|model deployment|model monitoring|data drift|vllm|agentic ai|rag|llm)\b"),
        ("DevSecOps & Cloud Security", r"\b(devsecops|security|iam|service account|organization polic|cloud armor|vpc service controls|zero trust|waf|owasp|ssl|tls|prisma cloud|wiz|ionix|sast|dast|secret manager|oauth|oidc|jwt)\b"),
        ("Monitoring, Logging & Observability", r"\b(monitoring|logging|observability|prometheus|grafana|dynatrace|elk|elasticsearch|logstash|kibana|jaeger|opentelemetry|sli|slo|telemetry|alert)\b"),
        ("Programming & Scripting", r"\b(python|bash|shell script|go|javascript|coding|script|automation code)\b"),
        ("Databases & Data Services", r"\b(cloud sql|alloydb|bigquery|postgresql|mysql|mongodb|redis|database|datastore|data service)\b"),
        ("Messaging & Streaming", r"\b(pub/sub|kafka|rabbitmq|messag|streaming|event-driven|queue)\b"),
        ("Source Control & Artifact Management", r"\b(git|github|gitlab|bitbucket|artifact registry|jfrog|source control|artifact management)\b"),
        ("Backup & Disaster Recovery", r"\b(backup|disaster recovery|dr|restore|veeam|kasten|rpo|rto|business continuity)\b"),
        ("ITSM & Enterprise Tools", r"\b(servicenow|jira|itsm|change management|problem management)\b"),
        ("Platform Engineering & SRE", r"\b(platform engineering|sre|production support|on-call|incident|root cause|capacity planning|cost optimization|rightsizing|developer experience|reliability)\b"),
        ("Hybrid & Multi-Cloud", r"\b(hybrid|multi-cloud|multicloud|aws|azure|on-prem|multi-region)\b"),
        ("Application & API Technologies", r"\b(fastapi|spring boot|node\.js|react|rest api|microservice|application|api technolog)\b"),
    ]
    for topic, pattern in rules:
        if re.search(pattern, text):
            return topic
    return "Cloud Platforms & GCP Services"

def build_example(entry):
    text = " ".join([
        entry.get("section") or "",
        entry.get("category") or "",
        entry.get("question") or "",
    ]).lower()
    examples = [
        (r"\b(shared vpc|vpc peering|cloud router|bgp|ha vpn|interconnect|cloud nat|private service connect|psc)\b",
         "For example, place application projects in a Shared VPC, advertise only approved CIDRs through Cloud Router and BGP, and verify the path with Connectivity Tests and VPC Flow Logs before changing production routes."),
        (r"\b(load balanc|health check|cloud armor)\b",
         "For example, expose a GKE service through a global HTTPS load balancer, use a `/healthz` backend health check, and apply a Cloud Armor rate-limit rule before allowing public traffic."),
        (r"\b(kubernetes|gke|pod|deployment|statefulset|daemonset|readiness|liveness|hpa|node pool)\b",
         "For example, deploy version v2 with a readiness probe and `maxUnavailable: 0`; Kubernetes adds a healthy v2 pod before terminating a v1 pod, preserving capacity during the rollout."),
        (r"\b(terraform|infrastructure as code|remote state|state lock|drift)\b",
         "For example, store Terraform state in a versioned GCS backend, run `terraform plan` in the pull request pipeline, require approval, and apply only the reviewed plan from the protected main branch."),
        (r"\b(jenkins|ci/cd|pipeline|gitops|argo cd|github actions|gitlab ci|canary|blue-green)\b",
         "For example, a Jenkins pipeline can run tests and security scans, publish an immutable image to Artifact Registry, update the GitOps repository, and let Argo CD promote it with an automatic rollback on failed health checks."),
        (r"\b(prometheus|grafana|opentelemetry|observability|monitoring|logging|sli|slo|error budget)\b",
         "For example, define availability as successful requests divided by total requests, set a 99.9% SLO, page on fast error-budget burn, and use traces plus correlated logs to identify the failing dependency."),
        (r"\b(iam|workload identity|secret|zero trust|waf|owasp|sast|dast|devsecops|security)\b",
         "For example, map a Kubernetes service account to a least-privilege Google service account with Workload Identity, keep credentials in Secret Manager, and block high-risk requests with Cloud Armor."),
        (r"\b(mlops|vertex ai|mlflow|kubeflow|model serving|model monitoring|data drift|llm|rag)\b",
         "For example, register a validated model in MLflow, deploy it to a canary endpoint, compare latency and prediction quality with the current model, and roll back automatically when drift or error thresholds are exceeded."),
        (r"\b(postgresql|mysql|mongodb|redis|cloud sql|alloydb|bigquery|database)\b",
         "For example, use a Cloud SQL read replica for read-heavy traffic, connection pooling to protect the primary, automated backups with point-in-time recovery, and a tested failover runbook."),
        (r"\b(kafka|rabbitmq|pub/sub|messag|streaming|queue)\b",
         "For example, publish orders with an idempotency key, retry transient consumer failures with exponential backoff, and route poison messages to a dead-letter topic for investigation."),
        (r"\b(backup|disaster recovery|rpo|rto|restore|failover)\b",
         "For example, target a 15-minute RPO and one-hour RTO, replicate backups to a separate region, run quarterly restore drills, and record the measured recovery time rather than assuming the design works."),
        (r"\b(python|bash|shell script|javascript|go programming|coding)\b",
         "For example, make the automation validate inputs, return non-zero exit codes on failure, emit structured logs, retry only transient errors, and include unit tests for success and failure paths."),
        (r"\b(rest api|fastapi|spring boot|node\.js|microservice|api)\b",
         "For example, expose an idempotent API endpoint with schema validation, OAuth-based authorization, request correlation IDs, timeouts, bounded retries, and a backward-compatible versioning policy."),
        (r"\b(incident|root cause|on-call|sre|production support)\b",
         "For example, stabilize the service first by rolling back the latest release, communicate impact and mitigation, build a timestamped incident timeline, and convert the RCA actions into owned, measurable follow-ups."),
        (r"\b(aws|azure|hybrid|multi-cloud|on-prem|cloud migration)\b",
         "For example, connect on-premises networks to GCP through redundant HA VPN tunnels, use dynamic BGP routing, test failover one tunnel at a time, and monitor route changes and packet loss."),
    ]
    for pattern, example in examples:
        if re.search(pattern, text):
            return example
    topic = classify_question(entry)
    return (
        f"For example, in a production {topic} review, document the requirement, compare at least "
        "two viable options, implement the safest option through automation, validate it with measurable "
        "success criteria, and retain a tested rollback path."
    )

topic_sections = {topic: [] for topic in TOPIC_SECTIONS}
for entry in entries:
    topic_sections[classify_question(entry)].append(entry)

doc = Document()

# ---- base style ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = rpr.makeelement(qn("w:rFonts"), {})
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Calibri")

# ---- title page ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI Mock Interview Question Bank")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Senior GCP DevOps / SRE / Platform Engineer / MLOps & LLMOps Engineer")
run.font.size = Pt(15)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f"{len(entries)} unique questions with answers  •  Generated {date.today().isoformat()}")
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.add_run(
    "This document combines every question and answer from the ai-mock-interviewer question bank: "
    "all questions are organized into the same 17 topic sections available in the application. "
    "Each answer retains its original source section and category for context."
)

doc.add_page_break()

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_qa(number, entry):
    q_para = doc.add_paragraph()
    q_para.paragraph_format.space_before = Pt(10)
    q_para.paragraph_format.space_after = Pt(2)
    question_text = entry["question"]
    if "\n" in question_text:
        lines = question_text.split("\n")
        q_run = q_para.add_run(f"Q{number}. {lines[0]}")
        q_run.bold = True
        q_run.font.size = Pt(11.5)
        for line in lines[1:]:
            code_run = q_para.add_run()
            code_run.add_break()
            code_run.text = line if line.strip() else " "
            code_run.font.name = "Consolas"
            code_run.font.size = Pt(9.5)
            code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    else:
        q_run = q_para.add_run(f"Q{number}. {question_text}")
        q_run.bold = True
        q_run.font.size = Pt(11.5)
    metadata = " • ".join(filter(None, [entry.get("category"), entry.get("questionType")]))
    if metadata:
        cat_para = doc.add_paragraph()
        cat_para.paragraph_format.space_before = Pt(0)
        cat_para.paragraph_format.space_after = Pt(2)
        cat_run = cat_para.add_run(metadata)
        cat_run.italic = True
        cat_run.font.size = Pt(9)
        cat_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    if entry.get("dateAdded"):
        date_para = doc.add_paragraph()
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(2)
        date_run = date_para.add_run(f"Added: {entry['dateAdded']}")
        date_run.italic = True
        date_run.font.size = Pt(8.5)
        date_run.font.color.rgb = RGBColor(0x2A, 0x78, 0xD6)
    a_para = doc.add_paragraph()
    a_para.paragraph_format.space_after = Pt(8)
    a_para.paragraph_format.left_indent = Inches(0.15)
    a_run = a_para.add_run("A: " + entry["answer"])
    a_run.font.size = Pt(10.5)
    example_para = doc.add_paragraph()
    example_para.paragraph_format.space_after = Pt(10)
    example_para.paragraph_format.left_indent = Inches(0.15)
    example_label = example_para.add_run("Example: ")
    example_label.bold = True
    example_label.font.size = Pt(10)
    example_run = example_para.add_run(build_example(entry))
    example_run.italic = True
    example_run.font.size = Pt(10)
    example_run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

# ---- Part 1: topic-organized question bank ----
add_heading("Part 1 — Topic-Wise Question Bank", level=1)
doc.add_paragraph(
    "Use the clickable document index to jump directly to any topic. Questions from fixed mock "
    "rounds and technical banks are consolidated here without duplication."
)

counter = 1
for topic_number, topic in enumerate(TOPIC_SECTIONS, 1):
    items = topic_sections[topic]
    add_heading(f"{topic_number}. {topic} ({len(items)} questions)", level=2)
    for entry in items:
        add_qa(counter, entry)
        counter += 1

# ---- Part 2: quick reference cheat sheet ----
doc.add_page_break()
add_heading("Part 2 — Quick Reference Cheat Sheet", level=1)
doc.add_paragraph(
    "Condensed keywords and reminders for last-minute review before an interview. "
    "The same content is available live during a practice session by opening the browser "
    "console and running cheatSheet('topic')."
)

for topic in cheat_sheets.values():
    add_heading(topic["label"], level=2)
    for point in topic["points"]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(point)
        run.font.size = Pt(10.5)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
print(f"Total questions in document: {len(entries)}")
